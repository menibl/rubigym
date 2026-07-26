from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "BALY-wellness-rules-review-he.docx"

GOLD = "C9A867"
DARK = "1D1F1B"
INK = "242621"
MUTED = "666961"
PALE = "F3F0E8"
GREEN = "E5F3EC"
RED = "FBE9E9"
AMBER = "FFF3D6"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D8D5CC", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph.alignment = align
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")
    return paragraph


def set_run(run, size=None, bold=None, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run._element.get_or_add_rPr().append(OxmlElement("w:rtl"))


def add_text(doc, text, style=None, bold=False, color=None, size=None, space_after=4):
    p = doc.add_paragraph(style=style)
    rtl(p)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text, level=0, color=None):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    rtl(p)
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.right_indent = Cm(0.45 + level * 0.45)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.paragraph_format.space_after = Pt(2.5)
    run = p.add_run(text)
    set_run(run, size=9.4, color=color or INK)
    return p


def add_checkbox(doc, text, status="review"):
    symbol = "☑" if status == "done" else "☐"
    color = "2F765B" if status == "done" else "9A6A16"
    p = doc.add_paragraph()
    rtl(p)
    p.paragraph_format.space_after = Pt(2.5)
    p.paragraph_format.right_indent = Cm(0.2)
    r1 = p.add_run(f"{symbol} ")
    set_run(r1, size=10, bold=True, color=color)
    r2 = p.add_run(text)
    set_run(r2, size=9.4, color=INK)


def heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    rtl(p)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(12 if level == 1 else 7)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_run(r, size=15 if level == 1 else 11.5, bold=True, color=GOLD if level == 1 else DARK)
    return p


def callout(doc, title, text, kind="info"):
    fill = {"info": PALE, "ok": GREEN, "warn": AMBER, "gap": RED}[kind]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, GOLD if kind in ("info", "warn") else ("6EAA8F" if kind == "ok" else "D99696"), "7")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    rtl(p)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + "\n")
    set_run(r, size=10, bold=True, color=DARK)
    r = p.add_run(text)
    set_run(r, size=9.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    bidi = OxmlElement("w:bidiVisual")
    bidi.set(qn("w:val"), "1")
    tbl_pr.append(bidi)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, DARK)
        set_cell_border(cell, "343730", "4")
        p = cell.paragraphs[0]
        rtl(p)
        r = p.add_run(header)
        set_run(r, size=8.6, bold=True, color=WHITE)
        if widths:
            cell.width = Cm(widths[idx])
    for row_i, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            set_cell_shading(cell, WHITE if row_i % 2 == 0 else "F8F7F3")
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            rtl(p)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run(r, size=8.4, color=INK)
            if widths:
                cell.width = Cm(widths[idx])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_page_number(paragraph):
    rtl(paragraph, WD_ALIGN_PARAGRAPH.CENTER)
    run = paragraph.add_run("BALY wellness  •  ")
    set_run(run, size=8, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)
section.header_distance = Inches(0.25)
section.footer_distance = Inches(0.28)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"]._element.rPr.rFonts.set(qn("w:cs"), "Arial")
styles["Normal"].font.size = Pt(9.5)
styles["Normal"].font.color.rgb = RGBColor.from_string(INK)
for style_name in ("List Bullet", "List Bullet 2"):
    styles[style_name].font.name = "Arial"
    styles[style_name]._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    styles[style_name].font.size = Pt(9.4)

header = section.header
p = header.paragraphs[0]
rtl(p)
r = p.add_run("BALY")
set_run(r, size=12, bold=True, color=DARK)
r = p.add_run(" wellness  |  מסמך כללי מערכת")
set_run(r, size=8.5, color=GOLD)
add_page_number(section.footer.paragraphs[0])

# Cover / executive summary
p = doc.add_paragraph()
rtl(p)
p.paragraph_format.space_before = Pt(20)
r = p.add_run("BALY")
set_run(r, size=18, bold=True, color=DARK)
r = p.add_run(" wellness")
set_run(r, size=12, color=GOLD)

add_text(doc, "חוקי הרשמה, מנויים וזכאות לאימונים", bold=True, color=DARK, size=25, space_after=4)
add_text(doc, "מסמך בדיקה ואישור לבעל המועדון", bold=True, color=GOLD, size=12, space_after=16)
add_text(doc, "תאריך בדיקה: 26.07.2026  |  גרסת מערכת: סביבת הדגמה", color=MUTED, size=9, space_after=12)

callout(
    doc,
    "מטרת המסמך",
    "להציג בשפה עסקית מה המערכת אוכפת כיום, מה נוסף בגרסה הנוכחית, ואילו החלטות עדיין דרושות לפני חיבור ספקי SMS וסליקה.",
    "info",
)
callout(
    doc,
    "סטטוס גרסת הבדיקה",
    "OTP קבוע: 1111. הודעת SMS ותשלום כרטיס אשראי הם סימולציה בלבד ואינם יוצרים חיוב או הודעה אמיתיים.",
    "warn",
)

heading(doc, "תקציר מנהלים", 1)
add_bullet(doc, "דף כניסה ציבורי חדש ממותג BALY wellness כולל הרשמה ו־LOGIN.")
add_bullet(doc, "הרשמה חדשה מתחילה באימות טלפון; לאחר האימות נאספים פרטי החשבון והצהרת הבריאות.")
add_bullet(doc, "כניסה אפשרית בשם משתמש/אימייל וסיסמה, או בטלפון ובקוד OTP.")
add_bullet(doc, "במובייל קיימת גישה ישירה לתוכנית אימונים, לפרופיל, לעדכון מנוי ולתשלום סימולציה.")
add_bullet(doc, "זכאות לאימון נבדקת לפי סטטוס תשלום, סוג מנוי, גיל, מין, קיבולת והגבלות האימון.")

heading(doc, "מקרא סטטוסים במסמך", 2)
add_table(
    doc,
    ["סימון", "משמעות"],
    [
        ["קיים", "החוק נאכף בקוד או זמין בממשק הבדיקה."],
        ["סימולציה", "התהליך עובד בממשק אך ללא ספק חיצוני או חיוב אמיתי."],
        ["להחלטה", "נדרשת החלטה עסקית לפני פיתוח או אכיפה."],
    ],
    [3.2, 13.0],
)

doc.add_page_break()

heading(doc, "1. הרשמה ופתיחת חשבון", 1)
add_table(
    doc,
    ["כלל", "התנהגות נוכחית", "סטטוס"],
    [
        ["מספר טלפון", "חובה להזין מספר תקין; מספר שכבר קיים במערכת אינו יכול להירשם שוב.", "קיים"],
        ["אימות SMS", "נשלח מסך קוד. בסביבת הבדיקה הקוד הקבוע הוא 1111.", "סימולציה"],
        ["שם משתמש", "חובה וייחודי; לא ניתן ליצור שני חשבונות באותו שם.", "קיים"],
        ["סיסמה", "חובה, לפחות 4 תווים.", "קיים"],
        ["גיל", "חשבון עצמאי ניתן לפתוח מגיל 16. צעיר יותר מצורף דרך חשבון משפחתי.", "קיים"],
        ["פרטים אישיים", "שם מלא, תאריך לידה ומין נדרשים להשלמת ההרשמה.", "קיים"],
        ["הצהרת בריאות", "חתימה חובה; תוקף לשנה. לאחר שנה נדרשת חתימה מחדש.", "קיים"],
        ["הסכם הצטרפות", "חתימה חובה על ההסכם, התקנון, מדיניות הביטולים והפרטיות.", "קיים"],
        ["מנוי התחלתי", "נוצר חשבון מתאמן במסלול קבוצתי חודשי ובסטטוס חוב עד לתשלום.", "קיים"],
    ],
    [3.3, 10.4, 2.5],
)

heading(doc, "2. כניסה לחשבון", 1)
add_bullet(doc, "שם משתמש או אימייל + סיסמה: המערכת מאמתת התאמה מלאה לחשבון קיים.")
add_bullet(doc, "טלפון + SMS: המערכת בודקת שמספר הטלפון קיים ורק אז מאפשרת הזנת OTP.")
add_bullet(doc, "בסביבת הבדיקה בלבד: OTP תקין הוא 1111.")
add_bullet(doc, "לאחר כניסה המשתמש מועבר ללוח המתאים לתפקידו: מנהל, מאמן או מתאמן.")

heading(doc, "3. תנאי סף לכל הרשמה לאימון", 1)
add_table(
    doc,
    ["בדיקה", "תנאי מעבר"],
    [
        ["משתמש", "המשתמש הוא מתאמן מחובר וחשבון המשתמש קיים."],
        ["הצהרת בריאות", "חתומה ובתוקף. הצהרה חסרה או שחלפה שנה מחתימתה חוסמת את כל האימונים."],
        ["תשלום", "מנוי פעיל; או אישור תשלום אופליין ממנהל; או מנוי משפחתי ששולם בידי המשלם הראשי."],
        ["תוקף מנוי", "תאריך התפוגה נבדק בכל הרשמה; מנוי שפג חוסם הרשמה גם אם נשמר סטטוס ACTIVE."],
        ["הקפאה", "מנוי מוקפא אינו מאפשר הרשמה עד לביטול ההקפאה."],
        ["סוג מנוי", "סוג המנוי הראשי והמסלולים המשניים נבדקים מול סוג האימון."],
        ["מגבלות אימון", "מין, גיל ורשימת סוגי המנוי המותרים נבדקים לפי הגדרת האימון."],
        ["חפיפת שעות", "לא ניתן להירשם לשני אימונים או למשבצת Open Gym חופפת."],
        ["מקום פנוי", "אם יש מקום — הרשמה; אם מלא — כניסה לרשימת המתנה."],
    ],
    [4.0, 12.2],
)

callout(
    doc,
    "כלל בריאות שאושר",
    "הצהרת בריאות לא חתומה או הצהרה שחלפה שנה ממועד חתימתה חוסמת הרשמה לאימון קבוצתי, אישי ו־Open Gym. ניתן לחתום מחדש מתוך הפרופיל.",
    "ok",
)

doc.add_page_break()

heading(doc, "4. מפת זכאות לפי סוג מנוי", 1)
add_table(
    doc,
    ["סוג מנוי", "אימונים קבוצתיים", "Open Gym", "אימון אישי", "תוכנית אימון"],
    [
        ["קבוצתי חודשי", "כן", "כן — כלול במנוי", "רק כתוספת", "לפי בקשה/תוספת"],
        ["קבוצתי שנתי", "כן", "כן — כלול במנוי", "רק כתוספת", "לפי בקשה/תוספת"],
        ["פתוח חודשי", "רק אם הותר במפורש", "כן", "רק כתוספת", "כלול"],
        ["פתוח שנתי", "רק אם הותר במפורש", "כן", "רק כתוספת", "כלול"],
        ["כרטיסייה", "אם האימון מתיר ויש יתרה", "כן, עם יתרה", "לא", "לא"],
        ["אימון אישי", "לא לבדו", "לא לבדו", "כן", "לפי המאמן"],
        ["תוכנית אימון", "לא מעניק כניסה לבדו", "לא מעניק כניסה לבדו", "לא", "כן"],
        ["תוכנית תזונה", "לא מעניק כניסה", "לא מעניק כניסה", "לא", "לא"],
    ],
    [3.3, 3.1, 2.7, 2.8, 3.0],
)

heading(doc, "5. אימונים קבוצתיים", 1)
add_bullet(doc, "המערכת בודקת מנוי קבוצתי חודשי/שנתי, או התאמה מפורשת של האימון לרשימת סוגי מנוי מותרים.")
add_bullet(doc, "אם האימון מוגבל לנשים או לגברים, המין בחשבון חייב להתאים.")
add_bullet(doc, "אם הוגדר גיל מינימלי או מקסימלי, גיל המתאמן חייב להיכלל בטווח.")
add_bullet(doc, "כאשר המכסה מלאה, המשתמש נכנס לרשימת המתנה במקום לרשימת הרשומים.")
add_bullet(doc, "בכרטיסייה נדרש לפחות ניקוב אחד; בהרשמה יורד ניקוב ובביטול בזמן הוא מוחזר.")

heading(doc, "6. אימונים אישיים", 1)
add_bullet(doc, "נדרש מסלול PERSONAL_TRAINING כמנוי ראשי או משני.")
add_bullet(doc, "תמחור אימון אישי נשמר בפרופיל ויכול להשתנות בין מתאמנים.")
add_bullet(doc, "בביטול מאוחר חל חיוב מלא/ללא החזר ונרשמת נקודה שחורה.")

heading(doc, "7. Open Gym", 1)
add_bullet(doc, "מנוי קבוצתי חודשי ומנוי קבוצתי שנתי כוללים כניסה ל־Open Gym כחלק מהמנוי, ללא תשלום נוסף.")
add_bullet(doc, "נדרש סטטוס תשלום תקין.")
add_bullet(doc, "בכרטיסייה נדרשת יתרה חיובית ויורד ניקוב בהרשמה.")
add_bullet(doc, "כאשר המקום מלא, המשתמש מצורף לרשימת המתנה.")
callout(
    doc,
    "כלל שאושר",
    "הכניסה ל־Open Gym נאכפת לפי מפת הזכאות: קבוצתי חודשי/שנתי, פתוח חודשי/שנתי או כרטיסייה עם יתרה. נבדקים גם תשלום תקין והקפאת מנוי.",
    "ok",
)

doc.add_page_break()

heading(doc, "8. רשימת המתנה ועדיפות", 1)
add_bullet(doc, "סדר ההמתנה נקבע לפי זמן ההצטרפות בלבד — כל הקודם זוכה.")
add_bullet(doc, "ביטול מקום של נרשם מקדם אוטומטית את הראשון בתור.")
add_bullet(doc, "נקודות שחורות אינן משנות את מיקום המשתמש ברשימת ההמתנה.")
add_bullet(doc, "בקידום מהמתנה תישלח הודעת PUSH למשתמש כאשר PUSH פעיל במכשיר.")
callout(
    doc,
    "התראות PUSH",
    "המשתמש יכול להפעיל או לבטל PUSH ותזכורות אימון. מנהל יכול להפעיל PUSH עבור פניות שמופנות אליו. שליחה כשהיישום סגור דורשת ספק Push ו־Service Worker ב־Production.",
    "info",
)

heading(doc, "9. ביטולים, אי־הגעה ועונשים", 1)
add_table(
    doc,
    ["אירוע", "כלל נוכחי", "תוצאה"],
    [
        ["ביטול קבוצתי בזמן", "לפני חלון הביטול שמוגדר במערכת (ברירת מחדל 2 שעות).", "ללא נקודה; ניקוב מוחזר."],
        ["ביטול קבוצתי מאוחר", "בתוך חלון הביטול.", "נקודה שחורה; ניקוב אינו מוחזר."],
        ["ביטול אימון אישי", "בתוך חלון הביטול האחיד של 2 שעות.", "נקודה שחורה וחיוב מלא."],
        ["אי־הגעה", "נרשם לאימון אך לא ביצע צ'ק־אין.", "נקודה שחורה אוטומטית בסריקת הסיום."],
        ["פינוי מקום", "משתמש מבטל הרשמה.", "הראשון ברשימת ההמתנה מקודם."],
    ],
    [3.4, 8.2, 4.0],
)

doc.add_page_break()

heading(doc, "10. מנוי, תשלום ופרופיל", 1)
add_bullet(doc, "המשתמש יכול לערוך שם, שם משתמש, טלפון, תאריך לידה וסיסמה.")
add_bullet(doc, "המשתמש יכול לעבור בין מנוי קבוצתי חודשי לשנתי.")
add_bullet(doc, "תשלום/חידוש מנוי יוצר רשומת תשלום סימולציה, מפעיל את המנוי ומעדכן תוקף.")
add_bullet(doc, "ניתן להקפיא ולהפשיר מנוי; ביטול מוקדם של מנוי שנתי כולל קנס סימולציה.")
add_bullet(doc, "כרטיס דיגיטלי מאפשר סימולציית צ'ק־אין לאימון שאליו המשתמש רשום.")
add_table(
    doc,
    ["נושא", "כלל מערכת שנדרש לסליקה"],
    [
        ["מחירים ומע״מ", "מחירי המסלולים מוצגים לפני אישור; המחיר לצרכן כולל מע״מ ומצוין גם בקבלה."],
        ["חיוב מחזורי", "מנוי חודשי מתחדש מדי חודש. מנוי שנתי הוא התחייבות ל־12 חודשים עם חיוב חודשי."],
        ["אישור תשלום", "הפעלה רק לאחר אישור ספק הסליקה; נשמרים סכום, מסלול, מועד ואסמכתה."],
        ["קבלה", "קבלה דיגיטלית נשלחת למשתמש ונשמרת בהיסטוריית התשלומים."],
        ["כשל חיוב", "המנוי עובר לסטטוס חוב, הרשמה נחסמת ונשלחת התראת PUSH עם אפשרות לתשלום חוזר."],
        ["ביטול חודשי", "הביטול נכנס לתוקף בסוף התקופה שכבר שולמה; אין חיוב נוסף לאחר מכן."],
        ["ביטול שנתי", "ביטול מוקדם כפוף לקנס היציאה המוגדר במערכת ולתיעוד אישור המשתמש."],
        ["החזר", "החזר מבוצע רק לאחר אישור מנהל ונרשם כתנועת REFUNDED המקושרת לעסקה המקורית."],
        ["אבטחת מידע", "אין לשמור מספר כרטיס מלא או CVV; נשמר רק מזהה מאובטח מספק הסליקה."],
    ],
    [3.5, 12.7],
)
callout(
    doc,
    "הערת Production",
    "כעת התשלום הוא סימולציה. חיבור סליקה אמיתי מחייב ספק סליקה, חשבוניות/קבלות, Webhooks לאישור ולכשל חיוב, מדיניות החזר מאושרת ובדיקת ייעוץ משפטי/חשבונאי.",
    "warn",
)

doc.add_page_break()

heading(doc, "11. החלטות שאושרו", 1)
add_checkbox(doc, "פתיחת חשבון עצמאי מגיל 16; צעירים יותר מצורפים באמצעות חשבון משפחתי.", "done")
add_checkbox(doc, "הצהרת בריאות בתוקף לשנה בלבד; לאחר שנה נדרשת חתימה מחדש.", "done")
add_checkbox(doc, "תאריך תפוגת המנוי נבדק בכל הרשמה בנוסף לסטטוס המנוי.", "done")
add_checkbox(doc, "אסור להירשם לשני אימונים או משבצות חופפים.", "done")
add_checkbox(doc, "אין מגבלת כמות אימונים ביום או בשבוע, בכפוף לזכאות ולתפוסה.", "done")
add_checkbox(doc, "כללי Open Gym נקבעו לפי מפת הזכאות בסעיף 4.", "done")
add_checkbox(doc, "רשימת ההמתנה פועלת לפי זמן ההצטרפות — כל הקודם זוכה.", "done")
add_checkbox(doc, "PUSH כולל קידום מהמתנה ותזכורות אימון, עם אפשרות הפעלה וביטול.", "done")
add_checkbox(doc, "מנהל יכול להפעיל PUSH עבור הודעות ופניות שמופנות אליו.", "done")
add_checkbox(doc, "חלון הביטול אחיד: שעתיים לכל סוגי האימון.", "done")
add_checkbox(doc, "אי־הגעה ממשיכה ליצור נקודה שחורה אוטומטית.", "done")
add_checkbox(doc, "כל אחד מבני המשפחה רשאי לנהל, לשנות או לבטל מסלולים בחשבון המשפחתי.", "done")
add_checkbox(doc, "חתימה על הסכם ההצטרפות, התקנון, הביטולים והפרטיות היא חלק חובה מההרשמה.", "done")

heading(doc, "12. בדיקות קבלה מומלצות", 1)
add_checkbox(doc, "הרשמה עם טלפון חדש + OTP 1111 יוצרת משתמש בסטטוס חוב.", "done")
add_checkbox(doc, "ניסיון הרשמה עם טלפון קיים או שם משתמש קיים נחסם.", "done")
add_checkbox(doc, "כניסה בשם משתמש/סיסמה עובדת; כניסה בטלפון/OTP עובדת.", "done")
add_checkbox(doc, "בנייד ניתן להגיע לתוכנית אימונים ולפרופיל מהתפריט התחתון.", "done")
add_checkbox(doc, "עריכת פרופיל נפתחת ומאפשרת שינוי פרטים וסיסמה.", "done")
add_checkbox(doc, "תשלום סימולציה מפעיל מנוי ונרשם בהיסטוריית התשלומים.", "done")
add_checkbox(doc, "מתאמן ללא תשלום אינו יכול להירשם לאימון קבוצתי.")
add_checkbox(doc, "הגבלת גיל/מין חוסמת משתמש שאינו מתאים.")
add_checkbox(doc, "אימון מלא מעביר לרשימת המתנה וביטול מקדם את הראשון.")
add_checkbox(doc, "ביטול מאוחר ואי־הגעה יוצרים נקודה שחורה לפי המדיניות.")
add_checkbox(doc, "כרטיסייה יורדת בהרשמה וחוזרת רק בביטול בזמן.")
add_checkbox(doc, "הצהרה חסרה/שפג תוקפה חוסמת אימונים וחתימה מחדש פותחת את ההרשמה.", "done")
add_checkbox(doc, "מנוי שפג תוקפו או אימון חופף נחסמים.", "done")
add_checkbox(doc, "סדר ההמתנה נשמר לפי זמן ההצטרפות.", "done")
add_checkbox(doc, "ניתן להפעיל ולבטל PUSH ותזכורות; למנהל קיימת העדפת PUSH לפניות.", "done")
add_checkbox(doc, "ההרשמה כוללת חתימה על הסכם המועדון.", "done")

callout(
    doc,
    "המלצת סיום",
    "ההחלטות העסקיות שסופקו עודכנו במערכת. לפני פרסום ציבורי נותר לחבר ספקי SMS, PUSH וסליקה, להשלים Service Worker ולבצע בדיקות אבטחה וייעוץ משפטי/חשבונאי להסכמים ולמדיניות התשלום.",
    "ok",
)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
